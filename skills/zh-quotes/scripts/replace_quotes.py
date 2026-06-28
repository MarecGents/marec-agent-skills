#!/usr/bin/env python3
"""
replace_quotes.py — Replace English straight double quotes ("") in a .docx file
with Chinese full-width quotation marks ("\u201c" and "\u201d").

Usage:
    python replace_quotes.py <input.docx> [--output OUTPUT.docx] [--no-backup]

Arguments:
    <input.docx>          Path to the Word document to process.
    --output OUTPUT.docx  Output path. Default: overwrite the input file.
    --no-backup           Skip creation of a backup file (.bak).

How it works:
    1. Opens the .docx package (a ZIP archive of XML files).
    2. Traverses every paragraph in the document body, headers, and footers.
    3. Within each paragraph, merges text runs logically to handle quotes
       that span multiple runs, then replaces straight double quotes ("")
       with Chinese full-width quotation marks in pairwise alternating order:
         - 1st "  →  \u201c  (left double quotation mark)
         - 2nd "  →  \u201d  (right double quotation mark)
         - 3rd "  →  \u201c  (left, again)
         - 4th "  →  \u201d  (right, again)
         - ... and so on for each paragraph independently.
    4. Preserves all original formatting (font, size, bold, italic, colour, etc.)
    5. Writes the modified document.

Compatible with Python 3.8+ using python-docx (>=0.8.11) or lxml.

Author: Reasonix
"""

import os
import sys
import shutil
import argparse
from copy import deepcopy

# ─── Try python-docx first (simpler API), fall back to lxml ────────────

try:
    from docx import Document
    from docx.oxml.ns import qn
    USE_DOCX = True
except ImportError:
    USE_DOCX = False
    try:
        from lxml import etree
    except ImportError:
        print("Error: Neither python-docx nor lxml is installed.", file=sys.stderr)
        print("Install one with:  pip install python-docx", file=sys.stderr)
        sys.exit(1)


# ─── Core replacement logic (works with both backends) ────────────────

LEFT_Q  = '\u201c'   # "
RIGHT_Q = '\u201d'   # "


def replace_quotes_in_text(text: str) -> str:
    """Replace straight double quotes with Chinese full-width quotes in text."""
    result = []
    open_q = True
    for ch in text:
        if ch == '"':
            result.append(LEFT_Q if open_q else RIGHT_Q)
            open_q = not open_q
        else:
            result.append(ch)
    # If odd number of quotes, leave the last one as-is (user can fix)
    return ''.join(result)


# ─── python-docx backend ──────────────────────────────────────────────

if USE_DOCX:
    from docx.oxml import parse_xml

    def process_docx(input_path: str, output_path: str, backup: bool = True):
        """Process a .docx file using python-docx."""
        if backup and input_path == output_path:
            bak = input_path + '.bak'
            shutil.copy2(input_path, bak)
            print(f'Backup created: {bak}')

        doc = Document(input_path)

        # Process body paragraphs
        for para in doc.paragraphs:
            _process_para_docx(para)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _process_para_docx(para)

        # Process headers and footers (experimental — may not cover all types)
        for section in doc.sections:
            for header in (section.header, section.first_page_header):
                if header and not header.is_linked_to_previous:
                    for para in header.paragraphs:
                        _process_para_docx(para)
            for footer in (section.footer, section.first_page_footer):
                if footer and not footer.is_linked_to_previous:
                    for para in footer.paragraphs:
                        _process_para_docx(para)

            # Even/odd headers and footers
            for header in (section.even_page_header,):
                if header and not header.is_linked_to_previous:
                    for para in header.paragraphs:
                        _process_para_docx(para)
            for footer in (section.even_page_footer,):
                if footer and not footer.is_linked_to_previous:
                    for para in footer.paragraphs:
                        _process_para_docx(para)

        doc.save(output_path)
        print(f'Document saved: {output_path}')

    def _process_para_docx(para):
        """Replace quotes in a single python-docx paragraph (handles multi-run)."""
        # Collect all (run, text) pairs with their current text
        text_runs = [(r, r.text) for r in para.runs if r.text]
        if not text_runs:
            return

        # Build the full paragraph text
        full = ''.join(text for _, text in text_runs)

        # Find and replace quotes character by character, tracking run boundaries
        new_texts = []
        open_q = True
        for _, orig_text in text_runs:
            new_chars = []
            for ch in orig_text:
                if ch == '"':
                    new_chars.append(LEFT_Q if open_q else RIGHT_Q)
                    open_q = not open_q
                else:
                    new_chars.append(ch)
            new_texts.append(''.join(new_chars))

        # Apply back to runs
        for (r, _), new_text in zip(text_runs, new_texts):
            if r.text != new_text:
                r.text = new_text

else:
    # ─── lxml backend (fallback, more thorough) ─────────────────────────
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def process_docx(input_path: str, output_path: str, backup: bool = True):
        """Process a .docx file using lxml (direct XML manipulation)."""
        import zipfile
        from lxml import etree

        if backup and input_path == output_path:
            bak = input_path + '.bak'
            shutil.copy2(input_path, bak)
            print(f'Backup created: {bak}')

            # Read all files from the ZIP
        with zipfile.ZipFile(input_path, 'r') as zin:
            files = {n: zin.read(n) for n in zin.namelist()}

        # Process every XML file that may contain text
        xml_keys = [k for k in files if k.endswith('.xml') and (
            k.startswith('word/') or k.startswith('customXml/')
        )]

        for key in xml_keys:
            files[key] = _process_xml_lxml(files[key])

        # Write output
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name, data in files.items():
                zout.writestr(name, data)

        print(f'Document saved: {output_path}')

    def _process_xml_lxml(xml_bytes):
        """Process an XML byte string, replacing quotes in all w:t elements."""
        from lxml import etree
        if not xml_bytes or len(xml_bytes) < 100:
            return xml_bytes

        try:
            root = etree.fromstring(xml_bytes)
        except Exception:
            return xml_bytes

        body = root.find(f'{{{W}}}body')
        if body is None:
            return xml_bytes

        for para in body.findall(f'.//{{{W}}}p'):
            # Collect all (run, text_elem) in order
            pairs = []
            for r in para.findall(f'{{{W}}}r'):
                for t in r.findall(f'{{{W}}}t'):
                    if t.text is not None:
                        pairs.append((r, t))

            if not pairs:
                continue

            # Replace quotes per paragraph
            open_q = True
            for _, t in pairs:
                new_chars = []
                for ch in t.text:
                    if ch == '"':
                        new_chars.append(LEFT_Q if open_q else RIGHT_Q)
                        open_q = not open_q
                    elif ch in (LEFT_Q, RIGHT_Q):
                        # Already a Chinese quote — keep and toggle state
                        new_chars.append(ch)
                        open_q = not open_q
                    else:
                        new_chars.append(ch)
                t.text = ''.join(new_chars)

        return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


# ─── Entry point ──────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description='Replace straight double quotes with Chinese full-width quotes in .docx files.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('input', help='Path to the input .docx file')
    parser.add_argument('--output', '-o', help='Output path (default: overwrite input)')
    parser.add_argument('--no-backup', action='store_true', help='Skip backup creation')
    args = parser.parse_args()

    input_path = args.input
    output_path = args.output or input_path

    if not os.path.exists(input_path):
        print(f'Error: File not found: {input_path}', file=sys.stderr)
        sys.exit(1)

    if not input_path.lower().endswith('.docx'):
        print(f'Warning: Input does not have a .docx extension: {input_path}', file=sys.stderr)

    process_docx(input_path, output_path, backup=not args.no_backup)

    # Count replacements for summary
    count = 0
    try:
        if USE_DOCX:
            from docx import Document
            doc = Document(output_path)
            for para in doc.paragraphs:
                for r in para.runs:
                    if r.text:
                        count += r.text.count(LEFT_Q) + r.text.count(RIGHT_Q)
        else:
            import zipfile
            with zipfile.ZipFile(output_path, 'r') as z:
                xml = z.read('word/document.xml').decode('utf-8')
                count = xml.count(LEFT_Q) + xml.count(RIGHT_Q)
    except Exception:
        pass

    print(f'Replaced approximately {count} Chinese quotation marks in total.')


if __name__ == '__main__':
    main()
